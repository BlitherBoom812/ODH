import http.server
import os
import re
import socketserver
import json
import sys
import traceback

from bs4 import BeautifulSoup
from docx import Document
import logging

# 配置日志信息
logging.basicConfig(filename='app.log', level=logging.INFO,
                    format='%(asctime)s - %(levelname)s - %(message)s')
logging.basicConfig(filename='app.log', level=logging.WARNING,
                    format='%(asctime)s - %(levelname)s - %(message)s')
logging.basicConfig(filename='app.log', level=logging.ERROR,
                    format='%(asctime)s - %(levelname)s - %(message)s')

def add_line_to_document(document, document_name, line_text):
    # 添加一个新的段落到文档末尾
    document.add_paragraph(line_text)

    # 保存文档
    document.save(document_name)


def match_trans(text):

    # 创建BeautifulSoup对象进行HTML解析
    soup = BeautifulSoup(text, 'html.parser')

    # 查找包含翻译信息的<span>元素
    translation_span = soup.find('span', class_='tran')

    if translation_span:

        pos = soup.find('span', class_='pos').text  # 获取词性
        eng_translation = translation_span.find('span', class_='eng_tran').text  # 获取英文翻译
        chn_translation = translation_span.find('span', class_='chn_tran').text  # 获取中文翻译

        return pos, eng_translation, chn_translation

    else:
        logging.warning("未找到翻译部分")
        return None, None, None

def save_docx(word="", phone="", pos="", eng_trans="", chn_trans=""):
    
    pos = str(pos).lower()

    logging.info("save notes:")
    logging.info(word)
    logging.info(phone)
    logging.info(chn_trans)

    # 使用示例
    # desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    desktop_path = '.'
    docx_name = "words.docx"
    final_data_path = os.path.join(desktop_path, docx_name)
    line_text = f"{word}    {phone}    {pos}. {chn_trans}"

    if not os.path.exists(final_data_path):
        Document().save(final_data_path)
    
    document = Document(final_data_path)  # 打开已有的文档或创建一个新的文档
    add_line_to_document(document, final_data_path, line_text)

    logging.info(f"save notes: {line_text}")

# 创建一个简单的请求处理程序
def handle_request(request_json):
    action = request_json['action']
    # set response
    if action == "version":
        message = {
            "result": ["docx"],
            "error": None
        }
    elif action == "deckNames":
        message = {
            "result": ["docx"],
            "error": None
        }
    elif action == "modelNames":
        message = {
            "result": ["docx"],
            "error": None
        }
    elif action == "modelFieldNames":
        message = {
            "result": ["front", "back", "phone"],
            "error": None
        }
    else:
        message = {
            "result": ["docx"],
            "error": None
        }
    # parse data
    if 'params' in request_json.keys():
        params = request_json['params']
        logging.info('params')
        if 'note' in params.keys():
            note = params['note']
            logging.info('note')
            if 'fields' in note.keys():
                fields = note['fields']
                logging.info('fields')
                if 'front' in fields.keys() and 'back' in fields.keys() and 'phone' in fields.keys():
                    logging.info('front, back and phone')
                    front = fields['front']
                    back = fields['back']
                    phone = fields['phone']
                    pos, eng_trans, chn_trans = match_trans(back)
                    save_docx(front, phone, pos, eng_trans, chn_trans)
                else:
                    logging.warning("failed th parse note")
            else:
                logging.warning("failed th parse note")
    return message


# 设置服务器的 IP 地址和端口
host = "localhost"
port = 8765

class MyRequestHandler(http.server.SimpleHTTPRequestHandler):
    def do_POST(self):
        # 检查请求路径是否为 /
        if self.path == "/":
            # 获取请求的内容长度
            content_length = int(self.headers['Content-Length'])
            # 读取请求的数据
            post_data = self.rfile.read(content_length)
            # 将请求数据解码为字符串
            post_data_str = post_data.decode('utf-8')
            # 解析请求数据为 JSON 对象
            post_data_json = json.loads(post_data_str)

            # logging.info(json.dumps(post_data_json, indent=4, ensure_ascii=False))

            try:
                message = handle_request(post_data_json)
                self.send_response(200)
            except BaseException as e:
                logging.error(traceback.format_exc() + str(e))
                message = {
                    'result': ["docx"],
                    'error' : 'invalid json'
                }
                self.send_response(400)
            # 设置响应的 Content-Type 为 application/json
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            
            # 构建要返回的 JSON 响应
            response = message
            response_json = json.dumps(response)
            logging.info(f"response: {response_json}")
            
            # 发送响应数据
            self.wfile.write(response_json.encode('utf-8'))
        else:
            self.send_error(404)
def run_server():
    with socketserver.TCPServer((host, port), MyRequestHandler) as server:
        logging.info(f"Server started at http://{host}:{port}")
        # 启动服务器，一直运行直到停止
        server.serve_forever()

if __name__ == "__main__":
    # save_docx("hello", "verb", "greeting", "你好")
    # 创建服务器对象，并指定请求处理程序
    run_server()